from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm

import sys

from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QVBoxLayout, QFormLayout, QLineEdit, QPushButton, QFileDialog, QMessageBox


class DocxGenerator:
    def __init__(self, filename, data):
        self.filename = filename
        self.data = data

    def generate_docx(self):
        doc = Document()
        
        doc.add_paragraph('Document Generated with Data:')
        
        print(type(self.data.items()))

        #for key, value in self.data.items():
        #    doc.add_paragraph(f'{key}: {value}')# Create a new Document
        doc = Document()

        # Table 1: Contains a single cell with the text "GIORNALE DI CANTIERE - Nr. 049"
        table1 = doc.add_table(rows=1, cols=1)
        table1.style = 'Table Grid'
        cell1 = table1.cell(0, 0)
        paragraph = cell1.paragraphs[0]
        run = paragraph.add_run("GIORNALE DI CANTIERE - Nr." + self.data["NR. Giornale di Cantiere"])
        run.font.size = Pt(18)  # Large font size
        run.font.italic = True
        run.font.bold = True
        run.font.all_caps = True  # All caps
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align


        doc.add_paragraph()

        # Table 2: Contains four rows, each with a single cell
        table2 = doc.add_table(rows=4, cols=1)
        table2.style = 'Table Grid'
        # First row: Title in all caps and red color
        cell2_1 = table2.cell(0, 0)
        paragraph = cell2_1.paragraphs[0]
        run = paragraph.add_run("ATTIVITÀ DIURNE ")
        run.font.size = Pt(14)  # Default font size
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
        run.font.all_caps = True  # All caps
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align

        # Second row: "Opere: "
        cell2_2 = table2.cell(1, 0)
        paragraph = cell2_2.paragraphs[0]
        paragraph.add_run("Opere: " + self.data["Opere"])

        # Third row: "Data:"
        cell2_3 = table2.cell(2, 0)
        paragraph = cell2_3.paragraphs[0]
        paragraph.add_run("Data: " + self.data["Data"])

        # Fourth row: "Luogo:  Carreggiata:  No. operatori presenti: "
        cell2_4 = table2.cell(3, 0)
        paragraph = cell2_4.paragraphs[0]
        paragraph.add_run("Luogo: " + self.data["Luogo"] + "\tCarreggiata: " + self.data["Carreggiata"] + "\tNo. operatori presenti: " + self.data["No. operatori presenti"])


        doc.add_paragraph()

        # Table 1: Contains a single cell with the text "GIORNALE DI CANTIERE - Nr. 049"
        table1 = doc.add_table(rows=1, cols=1)
        table1.style = 'Table Grid'
        cell1 = table1.cell(0, 0)
        paragraph = cell1.paragraphs[0]
        run = paragraph.add_run("ATTIVITÀ -")
        run.font.size = Pt(12)  # Large font size
        run.font.all_caps = True  # All caps
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align

        doc.add_paragraph()

        # Create the third table with one row and one cell
        table4 = doc.add_table(rows=1, cols=1)
        table4.style = 'Table Grid'
        cell3 = table4.cell(0, 0)
        paragraph3 = cell3.paragraphs[0]

        # Adding the text content with different paragraphs
        # TODO change name of lavorazioni in corso.
        content = (
            "Direzione Lavori Locale - BSA: AFRY Svizzera SA, "+ self.data["Direzione Lavori Locale - BSA, operatori"] + "operatore, "+ self.data["Direzione Lavori Locale - BSA, veicoli"]+"veicolo\n\n"
            "DLL eseguita: sorveglianza e controllo del cantiere per tutte le lavorazioni in corso ("+self.data["DLL eseguita"] + ")\n\n"
            "Stato DLL: servizio costante programmato, sino al termine dei lavori.\n\n"
        )
        run3 = paragraph3.add_run(content)
        run3.font.size = Pt(11)
        paragraph3.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Center-align the text


        doc.add_paragraph()

        main_table = doc.add_table(rows=3, cols=1)
        main_table.style = 'Table Grid'

        # First row
        cell1 = main_table.cell(0, 0)
        paragraph1 = cell1.paragraphs[0]
        run1 = paragraph1.add_run(("Lotto 8130 - INEL_IAU\t\tDitta: Kummler + Matter EVT SA\n"
                                "Operai e mezzi d'opera impiegati dall'impresa: "))
        run1.font.size = Pt(11)  # Font size for emphasis
        run1.font.bold = True
        paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Center-align the text

        # Second row - Contains a sub-table
        cell2 = main_table.cell(1, 0)

        # Create the sub-table with 7 rows and 4 columns
        sub_table = cell2.add_table(rows=7, cols=4)
        sub_table.style = 'Table Grid'
        # Configure the first row with two columns
        sub_table.cell(0, 0).text = "Mano d'opera"
        sub_table.cell(0, 1).text = "Mezzi di cantiere"
        sub_table.cell(0, 2).text = ""
        sub_table.cell(0, 3).text = ""

        sub_table.cell(1, 0).text = "Tecnici"
        sub_table.cell(1, 1).text = "-"
        sub_table.cell(1, 2).text = "Camioncini"
        sub_table.cell(1, 3).text = "-"

        sub_table.cell(2, 0).text = "Capo Cantiere"
        sub_table.cell(2, 1).text = "-"
        sub_table.cell(2, 2).text = "Furgoni"
        sub_table.cell(2, 3).text = "1"

        sub_table.cell(3, 0).text = "Capo Squadra"
        sub_table.cell(3, 1).text = "1"
        sub_table.cell(3, 2).text = "Camion con gru"
        sub_table.cell(3, 3).text = "-"

        sub_table.cell(4, 0).text = "Operai qualificati"
        sub_table.cell(4, 1).text = "1"
        sub_table.cell(4, 2).text = "Navicelle PLE"
        sub_table.cell(4, 3).text = "-"

        sub_table.cell(5, 0).text = "Aiutanti"
        sub_table.cell(5, 1).text = "-"
        sub_table.cell(5, 2).text = "Sollevatore con cesta"
        sub_table.cell(5, 3).text = "-"

        sub_table.cell(6, 0).text = "Autisti"
        sub_table.cell(6, 1).text = "-"
        sub_table.cell(6, 2).text = "Rimorchi"
        sub_table.cell(6, 3).text = "-"


        # Third row
        cell3 = main_table.cell(2, 0)
        paragraph3 = cell3.paragraphs[0]
        run3 = paragraph3.add_run("Interventi e operazioni: ")
        run3.font.size = Pt(11)  # Font size for emphasis
        run3.font.bold = True
        cell3.add_paragraph("Eseguito lo scollegamento dell'alimentazione dell'impianto illuminazione del sottopasso Ronco nell'armadio di distribuzione all'interno del sottopasso.", style="List Bullet")
        cell3.add_paragraph("Eseguito lo smontaggio di N°4 corpi luminosi a soffitto all'interno del sottopasso. ", style="List Bullet")
        #run4 = cell3.paragraphs[0].add_run("Eseguito lo scollegamento dell'alimentazione dell'impianto illuminazione del sottopasso Ronco nell'armadio di distribuzione all'interno del sottopasso.", style="List Bullet")
        #run4 = cell3.paragraphs[0].add_run("Eseguito lo smontaggio di N°4 corpi luminosi a soffitto all'interno del sottopasso. ", style="List Bullet")
        paragraph3.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Center-align the text

        cell3.add_paragraph()
        run4 = cell3.paragraphs[3].add_run("Note:")
        run4.font.size = Pt(11)  # Font size for emphasis
        run4.font.bold = True

        doc.add_paragraph()

        table_diversi = doc.add_table(rows=1,cols=1)
        table_diversi.style = 'Table Grid'
        main_cell = table_diversi.cell(0,0)
        paragraph = main_cell.paragraphs[0]
        run = paragraph.add_run("Diversi:")
        run.font.size = Pt(11)  # Font size for emphasis
        run.font.bold = True

        main_cell.add_paragraph("Eventi importanti, visite, sopralluoghi, riunioni, ecc.:", style="List Bullet")
        main_cell.add_paragraph("Termini e programma lavori:", style="List Bullet")
        main_cell.add_paragraph("Problemi tecnici e/o osservazioni:", style="List Bullet")
        main_cell.add_paragraph("Modifiche di progetto:", style="List Bullet")
        main_cell.add_paragraph("Impedimenti al traffico:", style="List Bullet")
        main_cell.add_paragraph("Problemi relativi alla sicurezza:", style="List Bullet")
        main_cell.add_paragraph("Problemi relativi alla qualità:", style="List Bullet")
        main_cell.add_paragraph("Unità Territoriale UT IV:", style="List Bullet")
        main_cell.add_paragraph("Impianti BSA:", style="List Bullet")
        main_cell.add_paragraph("Imprevisti:", style="List Bullet")
        main_cell.add_paragraph("Note DLL:", style="List Bullet")

        doc.add_paragraph()

        table_norme_sicurezza = doc.add_table(rows=1,cols=1)
        table_norme_sicurezza.style = 'Table Grid'
        main_cell = table_norme_sicurezza.cell(0,0)
        paragraph = main_cell.paragraphs[0]
        run = paragraph.add_run("Norme e Sicurezza:")
        run.font.size = Pt(11)  # Font size for emphasis
        run.font.bold = True

        main_cell.add_paragraph("", style="List Bullet")


        doc.add_paragraph()

        table_elenco_constatazione = doc.add_table(rows=1,cols=1)
        table_elenco_constatazione.style = 'Table Grid'
        main_cell = table_elenco_constatazione.cell(0,0)
        paragraph = main_cell.paragraphs[0]
        run = paragraph.add_run("Elenco constatazione e stato dei danni agli equipaggiamenti BSA:")
        run.font.size = Pt(11)  # Font size for emphasis
        run.font.bold = True

        main_cell.add_paragraph("Elenco danneggiamenti", style="List Bullet")

        doc.add_paragraph()


        # Table 1: Contains a single cell with the text "GIORNALE DI CANTIERE - Nr. 049"
        table1 = doc.add_table(rows=1, cols=1)
        table1.style = 'Table Grid'
        cell1 = table1.cell(0, 0)
        paragraph = cell1.paragraphs[0]
        run = paragraph.add_run("Galleria Fotografica")
        run.font.size = Pt(12)  # Large font size
        run.font.all_caps = True  # All caps
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align


        doc.add_picture("c:/Users/rbr598/Github/doc-gen/mysite/python-docx/cat.jpg", width=Cm(10))

        doc.save(self.filename)
        return self.filename




class FormWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle('Form to DOCX Generator')
        
        # Create a central widget and set layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout()
        central_widget.setLayout(layout)
        
        # Create a form layout
        form_layout = QFormLayout()
        self.fields = {}
        
        # Add fields to the form
        self.fields['NR. Giornale di Cantiere'] = QLineEdit()
        self.fields['Opere'] = QLineEdit()
        self.fields['Data'] = QLineEdit()
        self.fields['Luogo'] = QLineEdit()
        self.fields['Carreggiata'] = QLineEdit()
        self.fields['No. operatori presenti'] = QLineEdit()
        self.fields['Direzione Lavori Locale - BSA, operatori'] = QLineEdit()
        self.fields['Direzione Lavori Locale - BSA, veicoli'] = QLineEdit()
        self.fields['DLL eseguita'] = QLineEdit()


        #self.fields['Lotto'] = QLineEdit()
        #self.fields['Ditta'] = QLineEdit()

        #self.fields['Tecnici'] = QLineEdit()
        #self.fields['Capo Cantiere'] = QLineEdit()
        #self.fields['Capo Squadra'] = QLineEdit()
        #self.fields['Operai Qualificati'] = QLineEdit()
        #self.fields['Aiutanti'] = QLineEdit()
        #self.fields['Autisti'] = QLineEdit()
        #self.fields['Camioncini'] = QLineEdit()
        #self.fields['Furgoni'] = QLineEdit()
        #self.fields['Camion con gru'] = QLineEdit()
        #self.fields['Navicelle PLE'] = QLineEdit()
        #self.fields['Sollevatore con cesta'] = QLineEdit()
        #self.fields['Rimorchi'] = QLineEdit()




        for key, widget in self.fields.items():
            form_layout.addRow(key, widget)
        
        layout.addLayout(form_layout)
        
        # Add a submit button
        submit_button = QPushButton('Generate DOCX')
        submit_button.clicked.connect(self.generate_docx)
        layout.addWidget(submit_button)
    
    def generate_docx(self):
        # Gather data from the form
        data = {key: widget.text() for key, widget in self.fields.items()}
        
        # Open a file dialog to select the save location
        filename, _ = QFileDialog.getSaveFileName(self, 'Save DOCX File', '', 'DOCX Files (*.docx)')
        
        if filename:
            try:
                # Generate the DOCX file
                generator = DocxGenerator(filename, data)
                generator.generate_docx()
                
                # Notify user of success
                QMessageBox.information(self, 'Success', f'DOCX file saved to {filename}')
            except Exception as e:
                QMessageBox.critical(self, 'Error', f'An error occurred: {e}')
        else:
            QMessageBox.warning(self, 'Warning', 'No file selected.')

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = FormWindow()
    window.show()
    sys.exit(app.exec_())